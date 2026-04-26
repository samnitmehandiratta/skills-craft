"""
Extracts plain text from PDF, DOCX, or image files.
PDF: pdfplumber (pure Python, no system deps)
DOCX: python-docx
Image: Pillow preprocessing + gemini-flash-1.5 vision via OpenRouter
"""
import io
import json
import os
from typing import Literal

import pdfplumber
import docx
from PIL import Image
from modules.validation.bedrock_client import vision_ocr

MAX_CHARS = 40_000
DocumentType = Literal["pdf", "docx", "image"]


def detect_document_type(filename: str, content_type: str) -> DocumentType:
    name = (filename or "").lower()
    ct = (content_type or "").lower()

    if name.endswith(".pdf") or "pdf" in ct:
        return "pdf"
    if name.endswith(".docx") or name.endswith(".doc") or "wordprocessingml" in ct or "msword" in ct:
        return "docx"
    if any(name.endswith(ext) for ext in (".jpg", ".jpeg", ".png", ".webp", ".tiff", ".bmp")) or ct.startswith("image/"):
        return "image"

    raise ValueError(f"Unsupported file type: {filename} ({content_type}). Supported: PDF, DOCX, JPG, PNG, WEBP.")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        pages = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        return "\n\n".join(pages)[:MAX_CHARS]
    except Exception:
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(file_bytes))
        parts = []

        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)[:MAX_CHARS]
    except Exception:
        return ""


def extract_text_from_image(file_bytes: bytes) -> str:
    try:
        img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        if max(img.size) > 2048:
            ratio = 2048 / max(img.size)
            img = img.resize((int(img.size[0] * ratio), int(img.size[1] * ratio)), Image.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return vision_ocr(buf.getvalue())[:MAX_CHARS]
    except Exception:
        return ""


def _pdf_to_images(file_bytes: bytes) -> list[bytes]:
    """Render each PDF page to PNG bytes using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        images = []
        for page in doc:
            mat = fitz.Matrix(2, 2)  # 2x zoom → ~150dpi
            pix = page.get_pixmap(matrix=mat)
            images.append(pix.tobytes("png"))
        doc.close()
        return images
    except Exception:
        return []


def parse_document(file_bytes: bytes, filename: str, content_type: str) -> dict:
    """
    Top-level entry point for the upload route.
    For scanned PDFs (< 50 chars from pdfplumber), falls back to vision OCR page by page.

    Returns:
        {text, doc_type, char_count, extraction_ok, warning}
    """
    doc_type = detect_document_type(filename, content_type)

    warning = None
    if doc_type == "pdf":
        text = extract_text_from_pdf(file_bytes)
        if len(text) < 50:
            # Scanned/image-based PDF — render pages and run vision OCR
            pages = _pdf_to_images(file_bytes)
            if pages:
                page_texts = [extract_text_from_image(p) for p in pages]
                text = "\n\n".join(t for t in page_texts if t)
            if not text:
                warning = "Could not extract text from this PDF. It may be password-protected or corrupted."
    elif doc_type == "docx":
        text = extract_text_from_docx(file_bytes)
    else:
        text = extract_text_from_image(file_bytes)
        if not text:
            warning = "Image OCR failed. Please try uploading a PDF or DOCX instead."

    return {
        "text": text,
        "doc_type": doc_type,
        "char_count": len(text),
        "extraction_ok": len(text) > 50,
        "warning": warning,
    }
